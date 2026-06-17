# -*- coding: utf-8 -*-
# pip install python-docx

"""
Pandoc docx post-process filter.

This filter currently:
- formats markdown tables and table captions with custom docx styles;
- converts markdown fenced code blocks to the text-block style;
- removes manually typed numeric prefixes from auto-numbered headings.

Extension rule:
- Keep `_process_reserved_extension_slot` as the final registered processor.
- When adding a new processor, add it before `_process_reserved_extension_slot`.
- After finishing that extension, leave or add another reserved extension slot.
"""

import argparse
import re
import sys
from collections.abc import Iterator
from pathlib import Path
from typing import Callable

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

DEBUG_MODE = False

STYLE_TABLE = "表格整体样式"
STYLE_HEADER = "表格表头样式"
STYLE_CELL_SHORT = "表格单元格正文样式 A"
STYLE_CELL_LONG = "表格单元格正文样式 B"
STYLE_CAPTION = "表题"
STYLE_TEXT_BLOCK = "文本块"
STYLE_TEXT_BLOCK_FALLBACK = "Block Text"
PANDOC_CODE_BLOCK_STYLES = {"Source Code", "Code Block"}
CAPTION_PATTERN = re.compile(r"^表(?:\s+|[:：]\s*)?(.*)$")
HEADING_STYLE_PATTERN = re.compile(r"^Heading [1-9]$")
MANUAL_HEADING_NUMBER_PATTERN = re.compile(
    r"^\s*\d+(?:[./．。]\d+|[./．。])*(?:[)）、])?\s+"
)
CHAR_THRESHOLD = 60
MAX_MERGE_COLS = 3
FORMAT_TABLE_ENABLED = False

BlockProcessor = Callable[[object, int, list[object], DocxDocument], None]


def _iter_block_items(parent: DocxDocument | _Cell) -> Iterator[Paragraph | Table]:
    """Yield paragraphs and tables under parent in document order."""
    if hasattr(parent, "element") and hasattr(parent.element, "body"):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("parent must be Document or Cell")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


BORDER_SINGLE = {"sz": "4", "val": "single", "color": "000000"}


def _set_cell_border(cell: _Cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        el = tcBorders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tcBorders.append(el)
        for key, val in BORDER_SINGLE.items():
            el.set(qn(f"w:{key}"), val)


def _iter_physical_cells(table: Table) -> Iterator[_Cell]:
    """Yield every physical w:tc in the table (including vMerge continuation cells)."""
    for row in table.rows:
        for tc in row._tr.iterchildren():
            if tc.tag == qn("w:tc"):
                yield _Cell(tc, row)


def _set_table_width_100(table: Table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), "5000")


def _cell_char_count(cell: _Cell) -> int:
    n = 0
    for p in cell.paragraphs:
        n += len(p.text)
    return n


def _cell_text(cell: _Cell) -> str:
    """Return concatenated plain text of a cell."""
    parts: list[str] = []
    for p in cell.paragraphs:
        if p.text:
            parts.append(p.text)
    return "".join(parts).strip()


def _merge_column_vertical(table: Table, col_idx: int) -> None:
    """
    Merge cells vertically in the given column.

    Rule:
    - A group starts at a row where the cell is non-empty.
    - The group extends downwards over subsequent rows where the cell is empty.
    - If a group spans more than one row, merge all cells in that group into
      the first row's cell.

    Uses table.cell(row, col) and stores cell refs before any merge so that
    after the first merge, row.cells[col] indexing does not shift.
    """
    row_count = len(table.rows)
    if row_count == 0:
        return
    try:
        table.cell(0, col_idx)
    except IndexError:
        return
    cells = [table.cell(r, col_idx) for r in range(row_count)]
    row = 0
    while row < row_count:
        text = _cell_text(cells[row])
        if not text:
            row += 1
            continue
        start = row
        end = row + 1
        while end < row_count:
            if _cell_text(cells[end]):
                break
            end += 1
        if end - start > 1:
            merged = cells[start]
            for merge_row in range(start + 1, end):
                merged = merged.merge(cells[merge_row])
        row = end


def _format_table_content(table: Table, max_merge_cols: int | None = None) -> None:
    """
    Format table content by vertically merging hierarchical columns.

    The first `max_merge_cols` columns are processed using `_merge_column_vertical`.
    This matches markdown tables where hierarchical structure is expressed by
    writing the group label only on the first row and leaving subsequent rows
    in that column empty.
    """
    if max_merge_cols is None:
        max_merge_cols = MAX_MERGE_COLS

    rows = table.rows
    if not rows:
        return
    col_count = len(rows[0].cells)
    if col_count == 0:
        return
    if len(rows) <= 2:
        return

    body_rows = list(range(1, len(rows)))
    first_col_cells = [table.cell(r, 0) for r in body_rows]
    has_first_col_blanks = any(_cell_text(c) == "" for c in first_col_cells)
    if not has_first_col_blanks:
        return
    for col_idx in range(min(max_merge_cols, col_count)):
        _merge_column_vertical(table, col_idx)


def _set_paragraph_style(
    paragraph: Paragraph,
    style_name: str,
    doc: DocxDocument,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> bool:
    try:
        if style_name in doc.styles:
            paragraph.style = doc.styles[style_name]
        else:
            paragraph.style = style_name
    except (KeyError, ValueError):
        return False
    if alignment is not None:
        paragraph.paragraph_format.alignment = alignment
    return True


def _set_paragraph_style_any(
    paragraph: Paragraph,
    style_names: tuple[str, ...],
    doc: DocxDocument,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> bool:
    for style_name in style_names:
        if _set_paragraph_style(paragraph, style_name, doc, alignment):
            return True
    return False


def _apply_caption_style(
    paragraph: Paragraph, caption_text: str, doc: DocxDocument
) -> None:
    if not _set_paragraph_style(
        paragraph, STYLE_CAPTION, doc, alignment=WD_ALIGN_PARAGRAPH.CENTER
    ):
        return
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = caption_text
    else:
        paragraph.add_run(caption_text)


def _extract_caption_text(text: str) -> str | None:
    """Return normalized caption text when the paragraph is a table caption."""
    match = CAPTION_PATTERN.match(text.strip())
    if match is None:
        return None
    caption_text = match.group(1).strip()
    if not caption_text:
        return None
    return caption_text


def _style_name(block: Paragraph) -> str:
    try:
        return block.style.name
    except (AttributeError, ValueError):
        return ""


def _is_pandoc_code_block(block: Paragraph) -> bool:
    return _style_name(block) in PANDOC_CODE_BLOCK_STYLES


def _is_heading(block: Paragraph) -> bool:
    return HEADING_STYLE_PATTERN.match(_style_name(block)) is not None


def _manual_heading_number_prefix_length(text: str) -> int:
    match = MANUAL_HEADING_NUMBER_PATTERN.match(text)
    if match is None:
        return 0
    return match.end()


def _delete_paragraph_prefix(paragraph: Paragraph, prefix_length: int) -> None:
    remaining = prefix_length
    for run in paragraph.runs:
        if remaining <= 0:
            return
        run_text = run.text
        if len(run_text) <= remaining:
            run.text = ""
            remaining -= len(run_text)
            continue
        run.text = run_text[remaining:]
        return


def _apply_text_block_style(paragraph: Paragraph, doc: DocxDocument) -> None:
    _set_paragraph_style_any(
        paragraph,
        (STYLE_TEXT_BLOCK, STYLE_TEXT_BLOCK_FALLBACK),
        doc,
    )


def _apply_table_styles(table: Table, doc: DocxDocument) -> None:
    rows = table.rows
    if not rows:
        return
    for cell in _iter_physical_cells(table):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    try:
        if STYLE_TABLE in doc.styles:
            table.style = doc.styles[STYLE_TABLE]
        else:
            table.style = STYLE_TABLE
    except (KeyError, ValueError):
        pass
    _set_table_width_100(table)
    for cell in _iter_physical_cells(table):
        _set_cell_border(cell)
    for cell in rows[0].cells:
        for p in cell.paragraphs:
            _set_paragraph_style(
                p, STYLE_HEADER, doc, alignment=WD_ALIGN_PARAGRAPH.CENTER
            )
    for row in rows[1:]:
        for cell in row.cells:
            if _cell_char_count(cell) > CHAR_THRESHOLD:
                style, alignment = STYLE_CELL_LONG, WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                style, alignment = STYLE_CELL_SHORT, WD_ALIGN_PARAGRAPH.CENTER
            for p in cell.paragraphs:
                _set_paragraph_style(p, style, doc, alignment=alignment)


def _process_table_and_caption(
    block: object,
    index: int,
    blocks: list[object],
    doc: DocxDocument,
) -> None:
    if isinstance(block, Paragraph) and index + 1 < len(blocks):
        next_block = blocks[index + 1]
        if isinstance(next_block, Table):
            text = (block.text or "").strip()
            caption_text = _extract_caption_text(text)
            if caption_text is not None:
                _apply_caption_style(block, caption_text, doc)
    if isinstance(block, Table):
        if FORMAT_TABLE_ENABLED:
            _format_table_content(block, MAX_MERGE_COLS)
        _apply_table_styles(block, doc)


def _process_text_block(
    block: object,
    index: int,
    blocks: list[object],
    doc: DocxDocument,
) -> None:
    del index, blocks
    if isinstance(block, Paragraph) and _is_pandoc_code_block(block):
        _apply_text_block_style(block, doc)


def _process_heading_manual_number(
    block: object,
    index: int,
    blocks: list[object],
    doc: DocxDocument,
) -> None:
    del index, blocks, doc
    if not isinstance(block, Paragraph) or not _is_heading(block):
        return
    prefix_length = _manual_heading_number_prefix_length(block.text)
    if prefix_length > 0:
        _delete_paragraph_prefix(block, prefix_length)


def _process_reserved_extension_slot(
    block: object,
    index: int,
    blocks: list[object],
    doc: DocxDocument,
) -> None:
    """
    Reserved extension slot.

    Do not remove this placeholder when no new extension is needed. If a future
    extension is implemented, register the new processor before this function
    and keep a reserved slot at the end of `_document_processors`.
    """
    del block, index, blocks, doc


def _document_processors() -> tuple[BlockProcessor, ...]:
    """
    Register document processors in execution order.

    Future extension point:
    - Add a new `_process_*` function with the BlockProcessor signature.
    - Register it before `_process_reserved_extension_slot`.
    - Keep `_process_reserved_extension_slot` as the final tuple item.
    """
    return (
        _process_table_and_caption,
        _process_text_block,
        _process_heading_manual_number,
        _process_reserved_extension_slot,
    )


def _process_blocks(
    parent: DocxDocument | _Cell,
    doc: DocxDocument,
    processors: tuple[BlockProcessor, ...],
) -> None:
    blocks = list(_iter_block_items(parent))
    for index, block in enumerate(blocks):
        for processor in processors:
            processor(block, index, blocks, doc)
        if isinstance(block, Table):
            for cell in _iter_physical_cells(block):
                _process_blocks(cell, doc, processors)


def process_document(doc: DocxDocument) -> None:
    _process_blocks(doc, doc, _document_processors())


def main() -> int:
    global CHAR_THRESHOLD, DEBUG_MODE, FORMAT_TABLE_ENABLED, MAX_MERGE_COLS
    parser = argparse.ArgumentParser(
        description="Apply custom post-process styles to a pandoc-generated docx.",
    )
    parser.add_argument(
        "input_docx",
        type=Path,
        help="Path to the docx file to process.",
    )
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=None,
        help="Output path. If omitted, overwrites the input file.",
    )
    parser.add_argument(
        "--threshold",
        type=int,
        default=CHAR_THRESHOLD,
        help=f"Character count threshold for long-cell style (default: {CHAR_THRESHOLD}).",
    )
    parser.add_argument(
        "--max-merge-cols",
        type=int,
        default=MAX_MERGE_COLS,
        help=(
            "Number of leading columns to vertically merge when a non-empty "
            "cell is followed by consecutive empty cells."
        ),
    )
    parser.add_argument("--log", action="store_true", help="Enable debug logging.")
    parser.add_argument(
        "--format-table",
        action="store_true",
        help=(
            "Format tables by vertically merging hierarchical columns using "
            "the first N leading columns (see --max-merge-cols)."
        ),
    )
    args = parser.parse_args()
    CHAR_THRESHOLD = args.threshold
    MAX_MERGE_COLS = max(0, args.max_merge_cols)
    DEBUG_MODE = args.log
    FORMAT_TABLE_ENABLED = args.format_table

    if not args.input_docx.is_file():
        print(f"Error: not a file: {args.input_docx}", file=sys.stderr)
        return 1

    out_path = args.output or args.input_docx
    try:
        doc = Document(str(args.input_docx))
        process_document(doc)
        doc.save(str(out_path))
    except FileNotFoundError:
        print(f"Error: file not found: {args.input_docx}", file=sys.stderr)
        return 1
    except Exception as e:
        if DEBUG_MODE:
            import traceback

            traceback.print_exc()
        print(f"Error: {e}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
